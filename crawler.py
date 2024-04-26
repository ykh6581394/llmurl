# -*- coding: utf-8 -*-
"""
Created on Tue Apr 23 06:47:24 2024

@author: 유국현
"""


from newspaper import Article
import requests
from bs4 import BeautifulSoup as bs
from docx import Document
import pandas as pd


## 1. Naver, Daum, Nate 뉴스 여부 확인하기
##   1-1. Naver뉴스인 경우 Sport, Enter, News 구분하여 HTML Parsing
##   1-2. Daum News인 경우 Error message Return
## 2. 사이트 Key 확인하기
##   2-1. URL Parsing 하여 Key 확인
##   2-2. URL에 따라 Error message Return 
## 3. Language 확인하기
##   3-1. 특정 사이트에 따라 


def articleLib(url, language):
    article = Article(url, language) # URL과 언어를 입력
    article.download()
    article.parse()
    title = article.title
    text = article.text
    date = article.publish_date
    
    return title, text, date
    
def domainSanitize(domain_finder):
    if domain_finder.startswith("www."):
        domain_finder_new = domain_finder[4:]
    else:
        domain_finder_new = domain_finder
        
    return domain_finder_new

def portalFinder(domain_finder_new):

    if "naver" in domain_finder_new:
        return "naver"
    elif "daum" in domain_finder_new:
        return "daum"
    elif "nate" in domain_finder_new:
        return "nate"           
    elif "zum" in domain_finder_new:
        return "zum" 
    else:
        return "media"

def languageFinder(domain_finder_new, eng_domain, jap_domain):
    if domain_finder_new in eng_domain:
        language = "en"
    elif domain_finder_new in jap_domain:
        language = "jp"
    else:
        language = "ko"
    return language

def reReturner(url_org, domain_finder_new, eng_domain, jap_domain):

    language = languageFinder(domain_finder_new, eng_domain, jap_domain)
    try:
        title, body, date = articleLib(url_org, language)
        if body == '':
            #print("Can't crawl the body")
            return "Bodyerror", title, body , date
        elif len(body)>10 and len(body)<100:
            #print("Read only Title")
            return "Bodyshorterror", title, body , date
        else:
            #print("Read Success")
            return "success", title, body , date
    except:
        #print("Site Blocked")
        return "Blockerror", "notitle", "nobody", "nodate"


def docxSaver(title, body, num):
    
    document = Document()
    document.add_heading(title, level=1)
    document.add_paragraph(body)
    
    document.save("./word/article_"+str(num)+'.docx')

def crawler(url_org, eng_domain, jap_domain, num):
    
    domain_finder = url_org.split("/")[2]
    domain_finder_new = domainSanitize(domain_finder)
    newscategory = portalFinder(domain_finder_new)

    if newscategory == "media":
        error0, title0, body0, date0 =  reReturner(url_org, domain_finder_new, eng_domain, jap_domain)
        docxSaver(title0, body0, num)
        return url_org, error0, title0, body0, date0
       
    elif newscategory == "naver":
        if domain_finder_new.startswith("entertain"):       
            try:
                page = requests.get(url_org)
                soup = bs(page.text, "html.parser")
                web = soup.find("div", "article_info").find("a").get("href")
                    
                domain_finder = web.split("/")[2]
                domain_finder_new = domainSanitize(domain_finder)
                error0, title0, body0, date0 = reReturner(web, domain_finder_new, eng_domain, jap_domain)
                docxSaver(title0, body0, num)
                return web, error0, title0, body0, date0
                
            except:
                return url_org, "Naver entertianments HTML Error" , None, None, None
                    
        
        elif domain_finder_new.startswith("sports"):       
            try:
                page = requests.get(url_org)
                soup = bs(page.text, "html.parser")
                web = soup.find("div", "info").find("a").get("href")
                        
                domain_finder = web.split("/")[2]
                domain_finder_new = domainSanitize(domain_finder)
                error0, title0, body0, date0 = reReturner(web, domain_finder_new, eng_domain, jap_domain)
                docxSaver(title0, body0, num)
                return web, error0, title0, body0, date0
          
            except:
                return url_org, "Naver sport HTML Error" , None, None, None
                 
        else:      
            try:
                page = requests.get(url_org)
                soup = bs(page.text, "html.parser")
                web = soup.find("div", "media_end_head_info_datestamp").find("a").get("href")
                        
                domain_finder = web.split("/")[2]
                domain_finder_new = domainSanitize(domain_finder)
                error0, title0, body0, date0 = reReturner(web, domain_finder_new, eng_domain, jap_domain)
                docxSaver(title0, body0, num)
                return web, error0, title0, body0, date0
            except:
                return url_org, "Naver news HTML Error" , None, None, None

    elif newscategory == "daum":
        return url_org, "Daum error", None, None, None
    
    elif newscategory == "nate":
        return url_org, "Daum error", None, None, None

    elif newscategory == "zum":
        return url_org, "Daum error", None, None, None

def main2(urls_all, eng_domain, jap_domain):
    error_all = []
    title_all = []
    body_all  = []
    for k in range(len(urls_all)):
        urls, error, title, body, date = crawler(urls_all[k], eng_domain, jap_domain, k)
        error_all.append(error)
        if title == "None":
            title.append("No Title")
        else:
            title_all.append(title)
        if body == "None":
            body_all.append("No Body")
        else:
            body_all.append(body)
    return error_all, title_all, body_all
    